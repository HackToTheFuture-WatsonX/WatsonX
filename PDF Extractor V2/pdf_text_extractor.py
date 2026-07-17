"""
PDF Text Extractor — Core Pipeline
====================================
Connects to Box Online via the Box SDK, finds PDF files in the configured
source folder, decrypts password-protected PDFs, extracts all text per page,
parses the structured background-check report layout, and exports results to:

  • Word  (.docx)  → Word Extracts/
  • Excel (.xlsx)  → CSV Extracts/
  • JSON  (.json)  → JSON File Extracts/

All configuration is read from config.json in the same directory.

Box Authentication
------------------
Currently uses a short-lived Developer Token (OAuth2 access_token) stored in
config.json. Developer Tokens expire every 60 minutes — generate a new one at:
  https://app.box.com/developers/console
For production / long-running use, replace the OAuth2 block in get_box_client()
with a JWTAuth or CCGAuth configuration from the Box SDK.

PDF Layout Notes
----------------
• This report uses a two-column layout: label on one line, value on the next.
• Status verdicts (Cleared / Not Cleared) come ONLY from the cover page and
  page 2 summary tables — never from the detail pages.
• Social Media Screening result appears on page 2 (index 1), so parse_summary_page()
  always receives pages[0] + pages[1] combined.
• Section headers are rendered as white text on coloured backgrounds — the
  parser identifies them by ALL-CAPS heading patterns, not by colour/font.
"""

import re
import json
import logging
from pathlib import Path
from datetime import datetime

import fitz                                       # PyMuPDF  — PDF parsing
from docx import Document                         # python-docx — Word export
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook                     # openpyxl — Excel export
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side
)
from openpyxl.utils import get_column_letter
from boxsdk import OAuth2, Client                 # boxsdk — Box API client


# ─────────────────────────────────────────────────────────────────────────────
# Directory / file paths
# All output folders are siblings of this script file.
# ─────────────────────────────────────────────────────────────────────────────
BASE_DIR     = Path(__file__).parent.resolve()    # folder containing this script
CONFIG_PATH  = BASE_DIR / "config.json"           # credentials + settings
WORD_OUT_DIR = BASE_DIR / "Word Extracts"         # .docx output root
CSV_OUT_DIR  = BASE_DIR / "CSV Extracts"          # .xlsx output root
JSON_OUT_DIR = BASE_DIR / "JSON File Extracts"    # .json output root
LOG_PATH     = BASE_DIR / "extractor.log"         # rolling activity log


# ─────────────────────────────────────────────────────────────────────────────
# Footer patterns — lines that repeat on every page as a page header/footer.
# They are stripped before any parsing to avoid polluting extracted fields.
# Add or adjust patterns here if the report template changes.
# ─────────────────────────────────────────────────────────────────────────────
_FOOTER_PATTERNS = [
    re.compile(r"^\d{1,2}F Marco Polo.*",  re.IGNORECASE),  # office address line
    re.compile(r"^Telephone:.*",           re.IGNORECASE),  # phone number line
    re.compile(r"^Email:.*",               re.IGNORECASE),  # email line
    re.compile(r"^Corpnet Global Corp\.",  re.IGNORECASE),  # company name line
]

# ─────────────────────────────────────────────────────────────────────────────
# Section heading patterns — ALL-CAPS headings that appear alone on a line.
# These are skipped by the key-value parser so they are not treated as field
# labels. Extend this list if new section types are added to the report.
# ─────────────────────────────────────────────────────────────────────────────
_SECTION_HEADINGS = re.compile(
    r"^(EMPLOYMENT CHECK|PROFESSIONAL REFERENCE CHECK|DATABASE CHECK|"
    r"ADVERSE MEDIA CHECK|GLOBAL SANCTIONS|BANKRUPTCY CHECK|"
    r"FINANCIAL/CREDIT CHECK|DIRECTORSHIP CHECK|CIVIL LITIGATION CHECK|"
    r"PROFESSIONAL LICENSE QUALIFICATION|SOCIAL MEDIA SCREENING|"
    r"VERIFICATION SUMMARY|REPORT TYPE|REPORT COLOR LEGEND)",
    re.IGNORECASE,
)


# ─────────────────────────────────────────────────────────────────────────────
# Activity logging — writes to extractor.log AND the console.
# ─────────────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    handlers=[
        logging.FileHandler(LOG_PATH, encoding="utf-8"),  # persistent log file
        logging.StreamHandler(),                           # console output
    ],
)
log = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Configuration loader
# ─────────────────────────────────────────────────────────────────────────────
def load_config() -> dict:
    """
    Read and return the contents of config.json as a dict.
    Raises FileNotFoundError if the file does not exist.
    """
    if not CONFIG_PATH.exists():
        raise FileNotFoundError(
            f"config.json not found at {CONFIG_PATH}. "
            "Please create it before running the extractor."
        )
    with open(CONFIG_PATH, "r", encoding="utf-8-sig") as fh:
        cfg = json.load(fh)
    log.info("Config loaded from %s", CONFIG_PATH)
    return cfg


# ─────────────────────────────────────────────────────────────────────────────
# Box API client factory
# ─────────────────────────────────────────────────────────────────────────────
def get_box_client(box_cfg: dict) -> Client:
    """
    Build and return an authenticated Box SDK Client using JWT (Service Account).

    Reads credentials from the file pointed to by box_cfg['jwt_config_file']
    (path relative to this script's directory or the web-app folder).
    JWT tokens are issued and rotated automatically — no manual refresh needed.
    """
    from boxsdk import JWTAuth

    # Resolve the JWT config file path
    jwt_filename = box_cfg.get("jwt_config_file", "box_jwt_config.json")
    candidates = [
        Path(__file__).parent / jwt_filename,                          # sibling of extractor
        Path(__file__).parent.parent / "WatsonX Challenge - Web" / jwt_filename,  # web-app folder
        Path(__file__).parent / ".." / "WatsonX Challenge - Web" / jwt_filename,
    ]
    jwt_path = next((p.resolve() for p in candidates if p.exists()), None)
    if jwt_path is None:
        raise FileNotFoundError(
            f"Box JWT config file '{jwt_filename}' not found.\n"
            f"Looked in: {[str(p.resolve()) for p in candidates]}"
        )

    log.info("Box: authenticating with JWT from %s", jwt_path)
    auth   = JWTAuth.from_settings_file(str(jwt_path))
    client = Client(auth)
    log.info("Box client ready (JWT / Service Account).")
    return client


# ─────────────────────────────────────────────────────────────────────────────
# Box folder scanner — finds all PDF files in a given Box folder
# ─────────────────────────────────────────────────────────────────────────────
def find_pdf_files_on_box(client: Client, folder_id: str,
                           search_subfolders: bool = True) -> list[dict]:
    """
    Return a list of {id, name} dicts for every .pdf file found in the Box
    folder identified by folder_id. Recursively scans sub-folders when
    search_subfolders is True.
    """
    results: list[dict] = []
    _scan_box_folder(client, folder_id, search_subfolders, results)
    log.info("Found %d PDF file(s) on Box.", len(results))
    return results


def _scan_box_folder(client: Client, folder_id: str,
                     recurse: bool, results: list[dict]) -> None:
    """
    Internal recursive helper for find_pdf_files_on_box.
    Appends matching PDF file metadata to the results list.
    Silently skips folders it cannot access (logs the error).
    """
    try:
        # Fetch up to 1 000 items — sufficient for typical report volumes
        folder = client.folder(folder_id).get_items(limit=1000)
    except Exception as exc:
        log.error("Could not list Box folder %s: %s", folder_id, exc)
        return

    for item in folder:
        if item.type == "file" and item.name.lower().endswith(".pdf"):
            # Collect file id and name; all other metadata is fetched on demand
            results.append({"id": item.id, "name": item.name})
            log.info("  Found PDF: %s (id=%s)", item.name, item.id)
        elif item.type == "folder" and recurse:
            log.info("  Scanning subfolder: %s (id=%s)", item.name, item.id)
            _scan_box_folder(client, item.id, recurse, results)


# ─────────────────────────────────────────────────────────────────────────────
# PDF downloader — streams file content into memory (no temp files on disk)
# ─────────────────────────────────────────────────────────────────────────────
def download_pdf_bytes(client: Client, file_id: str, file_name: str) -> bytes:
    """
    Download a Box file by file_id and return its raw bytes.
    The entire file is held in memory — suitable for the report sizes in use.
    For very large files (>100 MB) consider streaming to a temp file instead.
    """
    log.info("Downloading '%s' from Box...", file_name)
    file_content: bytes = client.file(file_id).content()
    log.info("Downloaded %d bytes.", len(file_content))
    return file_content


# ─────────────────────────────────────────────────────────────────────────────
# PDF decryption + text extraction
# Operates entirely on in-memory bytes — no files written to disk.
# ─────────────────────────────────────────────────────────────────────────────
def open_and_decrypt_pdf(pdf_bytes: bytes, file_name: str,
                          password: str) -> fitz.Document:
    """
    Open a PDF from raw bytes (encrypted or plain) and return a fitz.Document.
    If the PDF is password-protected, authenticates with the given password.
    Raises ValueError if the password is wrong.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        result = doc.authenticate(password)
        if result == 0:
            # Authentication failed — password is incorrect
            doc.close()
            raise ValueError(
                f"Incorrect password for '{file_name}'. "
                "Update 'pdf_password' in config.json."
            )
        log.info("'%s' decrypted successfully.", file_name)
    else:
        # File is not encrypted — no decryption needed
        log.info("'%s' is not password-protected.", file_name)

    return doc


def extract_text_by_page(doc: fitz.Document) -> list[str]:
    """
    Extract text from every page of a fitz.Document.
    Returns a list where index N is the raw text of page N+1.
    Uses PyMuPDF's plain-text mode which preserves reading order.

    NOTE: This extracts text as a flat string. Complex multi-column layouts
    may have lines interleaved — see strip_footer() and parse_kv_block() for
    how we handle the two-column format used by this report.
    """
    pages = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pages.append(page.get_text("text"))  # "text" mode: plain, preserves order
    return pages


# ─────────────────────────────────────────────────────────────────────────────
# Footer stripping
# Removes repeated header/footer lines that appear on every page, which would
# otherwise interfere with key-value parsing.
# ─────────────────────────────────────────────────────────────────────────────
def strip_footer(text: str) -> str:
    """
    Remove lines matching any pattern in _FOOTER_PATTERNS, then collapse
    runs of more than two consecutive blank lines down to a single blank line.
    Returns the cleaned text string.
    """
    cleaned = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            # Preserve blank lines (used as field separators)
            cleaned.append("")
            continue
        if any(p.match(stripped) for p in _FOOTER_PATTERNS):
            # This line is a footer/header repeat — discard it
            continue
        cleaned.append(line)

    # Collapse 3+ consecutive blank lines to a single blank line
    result = re.sub(r"\n{3,}", "\n\n", "\n".join(cleaned))
    return result.strip()


# ─────────────────────────────────────────────────────────────────────────────
# Status resolver
# The only three valid status values are "Cleared", "Not Cleared", and "--".
# "--" is the safe default — never guess when the text is ambiguous.
# ─────────────────────────────────────────────────────────────────────────────

# Keywords that indicate an explicitly positive (green) verification result
_CLEARED_KEYWORDS = re.compile(
    r"(verified\s*[–\-]\s*clear|^clear\b|cleared|no adverse|"
    r"no financial delinquency|no civil case|no directorship)",
    re.IGNORECASE,
)

# Keywords that indicate an explicitly negative (red) verification result.
# IMPORTANT: Only match phrases the PDF itself uses to flag a negative outcome.
# Factual statements like "Subject is not a licensed professional" are NOT red
# flags — the PDF would include "Not Verified" or "Red Flag" explicitly if
# it intended to flag the result as Not Cleared.
_NOT_CLEARED_KEYWORDS = re.compile(
    r"(not\s+verified\s*[–\-]|red\s+flag|unverified\s*[–\-]|"
    r"verified\s*[–\-]\s*(?!clear))",   # "Verified –" followed by anything except "clear"
    re.IGNORECASE,
)


def resolve_status(text: str) -> str:
    """
    Scan text for explicit status keywords and return one of:
      "Cleared"     — text contains a positive verification phrase
      "Not Cleared" — text contains a negative verification phrase
      "--"          — no conclusive phrase found; do NOT guess

    Always prefer "--" over an incorrect positive or negative judgement.
    """
    if not text or not text.strip():
        return "--"                             # empty input → unknown
    if _CLEARED_KEYWORDS.search(text):
        return "Cleared"
    if _NOT_CLEARED_KEYWORDS.search(text):
        return "Not Cleared"
    return "--"                                 # nothing conclusive found


# ─────────────────────────────────────────────────────────────────────────────
# Key-value block parser
# This report uses a two-line format for most fields:
#   Line N:   "Label:"          (label alone, colon at end)
#   Line N+1: "Value text"      (value on the next non-blank line)
# Some fields use inline format:  "Label: Value"
# Both formats are handled here.
# ─────────────────────────────────────────────────────────────────────────────

# Matches inline "Label: Value" (label and value on the same line)
_KV_INLINE_RE = re.compile(
    r"^([A-Za-z][A-Za-z0-9 _/()#'\u2019\-]{1,60}?):\s+(.+)$"
)

# Matches a label-only line ending with a colon (value follows on next line)
_KV_LABEL_RE = re.compile(
    r"^([A-Za-z][A-Za-z0-9 _/()#'\u2019\-]{1,60}?):\s*$"
)
# NOTE: \u2019 (curly apostrophe) is included in both patterns because some
# field labels in the PDF use a curly apostrophe, e.g. "Respondent's Name".
# Without it, those labels would fail to match and return empty values.


def parse_kv_block(text: str) -> dict:
    """
    Parse a block of text into a {label: value} dict.
    Processes lines in order; supports both inline and two-line formats.
    Section headings and bare value lines are ignored as potential keys.
    First occurrence wins — duplicate keys are not overwritten.
    """
    fields = {}
    lines  = [l.rstrip() for l in text.splitlines()]
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip known section headings — they are not field labels
        if _SECTION_HEADINGS.match(line):
            i += 1
            continue

        # ── Format A: "Label: Value" on the same line ────────────────────────
        m_inline = _KV_INLINE_RE.match(line)
        if m_inline:
            key   = m_inline.group(1).strip()
            value = m_inline.group(2).strip()
            if key not in fields:           # first occurrence wins
                fields[key] = value
            i += 1
            continue

        # ── Format B: "Label:" on one line, value on the next non-blank line ─
        m_label = _KV_LABEL_RE.match(line)
        if m_label:
            key = m_label.group(1).strip()

            # Look ahead to find the value line, skipping blank lines
            j           = i + 1
            value_parts = []
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1      # skip blank lines between label and value
                    continue
                # Stop if we've reached another label or a section heading
                if (_KV_LABEL_RE.match(next_line)
                        or _KV_INLINE_RE.match(next_line)
                        or _SECTION_HEADINGS.match(next_line)):
                    break
                value_parts.append(next_line)
                j += 1
                break           # only take the first value line

            value = " ".join(value_parts).strip()
            if key not in fields and value:
                fields[key] = value
            i = j
            continue

        i += 1   # line is neither a heading nor a label — skip it

    return fields


# ─────────────────────────────────────────────────────────────────────────────
# Summary page parser
# Parses the cover page (and page 2, which continues the summary table) to
# extract the report header fields and per-check verdict rows.
# ─────────────────────────────────────────────────────────────────────────────
def parse_summary_page(text: str) -> dict:
    """
    Parse the report summary section.

    Accepts the combined text of page 1 + page 2 so that checks that spill
    onto page 2 (e.g. Social Media Screening) are captured correctly.

    Returns a dict with:
      subject_name, overall_status, case_reference, case_received,
      package, delivery_date,
      employment_check_summary   — list of {employer, result, status, note}
      professional_reference_summary — list of {referee, result, status}
      database_check_summary     — list of {check, result, status}
    """
    clean  = strip_footer(text)
    fields = parse_kv_block(clean)

    # Subject name: find first standalone "Lastname, Firstname" pattern on its own line.
    # Falls back to the "Subject" KV field if the pattern is not found.
    name_match   = re.search(r"^([A-Z][a-z]+,\s+[A-Z][a-z]+)\s*$",
                              clean, re.MULTILINE)
    subject_name = name_match.group(1) if name_match else fields.get("Subject", "")

    # Overall report status: presence of the word "Cleared" on the cover page
    # is the report-level verdict as printed by the issuing company.
    overall_status = "Cleared" if re.search(r"\bCleared\b", clean) else "Not Cleared"

    # ── Employment and Reference check summary rows ───────────────────────────
    # The summary table lists each employer / referee on its own line followed
    # immediately by its result line (e.g. "Verified – Clear").
    # We scan between section heading lines to stay in the right section.
    emp_summary = []
    ref_summary = []
    lines       = clean.splitlines()
    in_emp      = False    # currently inside EMPLOYMENT CHECK section
    in_ref      = False    # currently inside PROFESSIONAL REFERENCE CHECK section

    for idx, line in enumerate(lines):
        stripped = line.strip()

        # Detect section boundaries to know which list to append to
        if re.match(r"EMPLOYMENT CHECK\s*$", stripped, re.IGNORECASE):
            in_emp, in_ref = True, False
            continue
        if re.match(r"PROFESSIONAL REFERENCE CHECK\s*$", stripped, re.IGNORECASE):
            in_emp, in_ref = False, True
            continue
        if re.match(r"DATABASE CHECK|SOCIAL MEDIA|OTHER CHECK", stripped, re.IGNORECASE):
            in_emp, in_ref = False, False   # exit both employment/reference sections
            continue

        if not stripped:
            continue    # skip blank lines inside sections

        if in_emp or in_ref:
            # Skip lines that are themselves result phrases (they belong to the
            # preceding name/company line, not a new entry)
            if re.match(r"(Verified|Not Verified|Clear)", stripped, re.IGNORECASE):
                continue

            # Look ahead up to 3 lines for the result phrase
            result = ""
            for j in range(idx + 1, min(idx + 4, len(lines))):
                next_stripped = lines[j].strip()
                if next_stripped:
                    if re.match(r"(Verified|Not Verified|Clear)",
                                next_stripped, re.IGNORECASE):
                        result = next_stripped
                    break   # stop at first non-blank line regardless

            if in_emp:
                # Exclude appendix notes that appear after the last employer row
                if re.match(r"\(See Appendix", stripped, re.IGNORECASE):
                    continue
                emp_summary.append({
                    "employer": stripped,
                    "result":   result,
                    "status":   resolve_status(result),
                    "note":     "",            # populated manually if needed
                })
            elif in_ref:
                ref_summary.append({
                    "referee": stripped,
                    "result":  result,
                    "status":  resolve_status(result),
                })

    # ── Database / Other check summary rows ──────────────────────────────────
    # These appear as a two-line block: check name on one line, result on next.
    db_summary = []
    _DB_NAMES  = [
        "Adverse Media Check",
        "Global Sanctions",
        "Bankruptcy Check",
        "Financial/Credit Check",
        "Directorship Check (DTI Only)",
        "Civil Litigation Check",
        "Professional License Qualification",
        "Social Media Screening",
    ]
    lines_list = clean.splitlines()
    for idx, line in enumerate(lines_list):
        stripped = line.strip()
        for check_name in _DB_NAMES:
            if stripped.lower() == check_name.lower():
                # Grab the result from the very next non-blank line
                for j in range(idx + 1, min(idx + 4, len(lines_list))):
                    next_line = lines_list[j].strip()
                    if next_line:
                        db_summary.append({
                            "check":  check_name,
                            "result": next_line,
                            "status": resolve_status(next_line),
                        })
                        break
                break   # matched — move to next line in outer loop

    return {
        "subject_name":                   subject_name,
        "overall_status":                 overall_status,
        "case_reference":                 fields.get("Case Reference No", ""),
        "case_received":                  fields.get("Case Receive Date", ""),
        "package":                        fields.get("Package", ""),
        "delivery_date":                  fields.get("Delivery Date", ""),
        "employment_check_summary":       emp_summary,
        "professional_reference_summary": ref_summary,
        "database_check_summary":         db_summary,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Employment check detail page parser
# ─────────────────────────────────────────────────────────────────────────────
def parse_employment_check(text: str, number: int,
                            summary_entry: dict | None = None) -> dict:
    """
    Parse a single EMPLOYMENT CHECK (N) detail page into a structured dict.

    summary_entry — the matching employer row from the cover-page summary,
    used to inject the correct result text and verification_status flag.
    Status is intentionally derived from the summary page, not the detail page,
    because the detail page does not contain explicit Cleared/Not Cleared text.
    """
    clean  = strip_footer(text)
    fields = parse_kv_block(clean)

    # Pull result and status from the summary page entry if available
    if summary_entry:
        result              = summary_entry.get("result", "")
        verification_status = summary_entry.get("status", resolve_status(clean))
    else:
        # No matching summary entry found — fall back to parsing the detail text
        # NOTE: this fallback is unlikely to produce a correct status because
        # the detail pages do not contain explicit verdict keywords.
        result              = ""
        verification_status = resolve_status(clean)

    return {
        "check_number":         number,
        "employer_name":        fields.get("Employer Name", ""),
        "company_address":      fields.get("Company Address", ""),
        "position_title":       fields.get("Position Title", ""),
        "dates_of_employment":  fields.get("Dates of Employment", ""),
        "status_of_employment": fields.get("Status of Employment", ""),
        "eligible_for_rehire":  fields.get("Eligible for Rehire?", ""),
        "reason_for_exit":      fields.get("Reason for Exit", ""),
        "respondents_name":     fields.get("Respondent\u2019s Name", ""),  # curly apostrophe
        "respondents_title":    fields.get("Respondent\u2019s Title", ""), # curly apostrophe
        "contact_details":      fields.get("Contact Details", ""),
        "verification_date":    fields.get("Verification Date", ""),
        "notes":                fields.get("Notes", ""),
        "result":               result,
        "verification_status":  verification_status,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Professional Reference check detail page parser
# ─────────────────────────────────────────────────────────────────────────────
def parse_reference_check(text: str, number: int,
                           summary_entry: dict | None = None) -> dict:
    """
    Parse a single PROFESSIONAL REFERENCE CHECK (N) detail page.

    summary_entry — the matching referee row from the cover-page summary,
    used to inject referee_name and verification_status.
    Check numbering starts at 1; check 1 maps to index 0 in ref_sum_list.
    """
    clean  = strip_footer(text)
    fields = parse_kv_block(clean)

    # Referee identity and verdict come from the summary page entry
    referee_name        = summary_entry.get("referee", "") if summary_entry else ""
    result              = summary_entry.get("result",  "") if summary_entry else ""
    verification_status = (
        summary_entry.get("status", resolve_status(clean))
        if summary_entry else resolve_status(clean)
    )

    # ── Extract interview Q&A pairs ───────────────────────────────────────────
    # Questions begin with interrogative words (What, Was, Did, …).
    # The answer immediately follows the last question-continuation line.
    qa_pairs = []
    lines    = [l.strip() for l in clean.splitlines() if l.strip()]
    i        = 0

    while i < len(lines):
        line = lines[i]

        if re.match(r"^(What|Was|Did|In |Would|How)", line, re.IGNORECASE):
            # Collect multi-line question text
            question_parts = [line]
            j = i + 1
            while (j < len(lines)
                   and not _KV_LABEL_RE.match(lines[j])
                   and not re.match(r"^(What|Was|Did|In |Would|How|Verif)",
                                    lines[j], re.IGNORECASE)):
                question_parts.append(lines[j])
                j += 1

            question = " ".join(question_parts)

            # The answer is the first non-question, non-verifier-signature line
            # after the question block.
            # "Verifier's Name and Designation" is a page signature placeholder,
            # never an actual answer — guard against it explicitly.
            answer = (
                lines[j]
                if j < len(lines)
                and not re.match(
                    r"^(What|Was|Did|In |Would|How|Verif)",
                    lines[j], re.IGNORECASE
                )
                else ""
            )

            qa_pairs.append({"question": question, "answer": answer})
            i = j + 1 if answer else j  # advance past the answer line if one was found
        else:
            i += 1

    return {
        "check_number":        number,
        "referee_name":        referee_name,
        "result":              result,
        "verifiers_name":      fields.get("Verifier\u2019s Name and Designation", ""),
        "verifiers_contact":   fields.get("Verifier\u2019s Contact Details", ""),
        "notes":               fields.get("Notes", ""),
        "verification_status": verification_status,
        "qa":                  qa_pairs,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Database / Other checks page parser
# ─────────────────────────────────────────────────────────────────────────────
def parse_other_checks(pages: list[str]) -> list[dict]:
    """
    Parse DATABASE CHECK / OTHER CHECK pages.

    Each named check section (ADVERSE MEDIA CHECK, GLOBAL SANCTIONS, etc.)
    becomes its own dict in the returned list with keys:
      check_name, source, result, status

    Only pages that do NOT contain an employment or reference check heading
    are processed, to avoid accidentally parsing those detail pages here.

    NOTE: The status and result fields here are later overwritten with values
    from the summary page by build_structured_json(), because the detail pages
    only describe the source/methodology — not the final verdict.
    """
    # Filter to only genuine database-check pages
    other_pages = []
    for p in pages:
        clean = strip_footer(p)
        # Exclude pages that are employment or reference check detail pages
        if not re.search(
            r"EMPLOYMENT CHECK\s*\(\d+\)|PROFESSIONAL REFERENCE CHECK\s*\(\d+\)",
            clean, re.IGNORECASE
        ):
            other_pages.append(clean)

    if not other_pages:
        return []   # no other-check pages found

    # Combine all other-check pages into one text block for section scanning
    combined = "\n".join(other_pages)

    # Regex to find each known check type as a standalone ALL-CAPS section heading
    section_re = re.compile(
        r"^(ADVERSE MEDIA CHECK|GLOBAL SANCTIONS|BANKRUPTCY CHECK|"
        r"FINANCIAL/CREDIT CHECK|DIRECTORSHIP CHECK|CIVIL LITIGATION CHECK|"
        r"PROFESSIONAL LICENSE QUALIFICATION|SOCIAL MEDIA SCREENING)\s*$",
        re.MULTILINE | re.IGNORECASE,
    )

    # Build a list of (char_offset, heading_title) for each found section
    sections   = []
    boundaries = [
        (m.start(), m.group(1).strip().upper())
        for m in section_re.finditer(combined)
    ]

    for idx, (start, title) in enumerate(boundaries):
        # Section body runs from after the heading to the start of the next heading
        end  = boundaries[idx + 1][0] if idx + 1 < len(boundaries) else len(combined)
        body = combined[start + len(title):end].strip()

        # Parse result lines and source lines from the section body
        source_lines = []
        result_lines = []
        in_source    = False    # flag: currently reading the Source block

        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                if in_source:
                    break       # blank line ends the source block
                continue

            if stripped.lower().startswith("source:"):
                # Begin collecting source text (may span multiple lines)
                in_source = True
                after = stripped[7:].strip()
                if after:
                    source_lines.append(after)
                continue

            if in_source:
                source_lines.append(stripped)
            else:
                result_lines.append(stripped)

        source = " ".join(source_lines).strip()
        result = " ".join(result_lines).strip()

        sections.append({
            "check_name": title.title(),                  # e.g. "Adverse Media Check"
            "source":     source,
            "result":     result if result else "(see source)",
            # Status resolved from the full section text; will be overwritten
            # by the summary page verdict in build_structured_json()
            "status":     resolve_status(combined[start:end]),
        })

    return sections


# ─────────────────────────────────────────────────────────────────────────────
# Main structured-JSON builder
# Orchestrates all page parsers and cross-references summary verdicts into
# the detail sections to produce one complete structured document per PDF.
# ─────────────────────────────────────────────────────────────────────────────
def build_structured_json(file_name: str, pages: list[str]) -> dict:
    """
    Route each PDF page to the appropriate parser and assemble the full
    structured output dict.

    Page routing:
      Page 0 (index 0)  — always the report summary / cover page
      Pages 1+          — routed by heading: employment, reference, or other
    """
    employment_pages  = {}   # {check_number (int): page_text}
    reference_pages   = {}   # {check_number (int): page_text}
    other_check_pages = []   # list of page_text for database/other checks

    # Route pages 1+ by their section heading
    for page_text in pages[1:]:
        clean = strip_footer(page_text)

        # Detect EMPLOYMENT CHECK (N) heading
        emp_m = re.search(r"EMPLOYMENT CHECK\s*\((\d+)\)", clean, re.IGNORECASE)
        # Detect PROFESSIONAL REFERENCE CHECK (N) heading
        ref_m = re.search(r"PROFESSIONAL REFERENCE CHECK\s*\((\d+)\)", clean, re.IGNORECASE)
        # Detect any database/other check heading
        db_m  = re.search(
            r"^(ADVERSE MEDIA CHECK|GLOBAL SANCTIONS|BANKRUPTCY CHECK|"
            r"FINANCIAL/CREDIT CHECK|DIRECTORSHIP CHECK|CIVIL LITIGATION CHECK|"
            r"PROFESSIONAL LICENSE QUALIFICATION|SOCIAL MEDIA SCREENING)\s*$",
            clean, re.IGNORECASE | re.MULTILINE,
        )

        if emp_m:
            employment_pages[int(emp_m.group(1))] = page_text
        elif ref_m:
            reference_pages[int(ref_m.group(1))] = page_text
        elif db_m:
            other_check_pages.append(page_text)
        # Pages that match none of the above (e.g. appendix pages) are silently ignored

    # ── Parse summary (cover + page 2 combined) ───────────────────────────────
    # Page 2 (index 1) contains the Social Media Screening result in this
    # report format, so we always merge it with page 1 for summary parsing.
    summary_text = pages[0] if pages else ""
    if len(pages) > 1:
        summary_text = summary_text + "\n" + pages[1]
    summary = parse_summary_page(summary_text) if pages else {}

    # ── Build lookup maps so detail parsers receive the correct summary entry ─
    # Employment: matched by employer name (case-insensitive substring)
    emp_sum_by_name = {
        e["employer"].lower(): e
        for e in summary.get("employment_check_summary", [])
    }
    # Reference: matched by ordinal position (check 1 → index 0, etc.)
    ref_sum_list = summary.get("professional_reference_summary", [])

    # ── Parse employment detail pages ─────────────────────────────────────────
    employment_checks = []
    for num, text in sorted(employment_pages.items()):
        clean    = strip_footer(text)
        emp_name = ""

        # Try inline "Employer Name: <value>" first
        for line in clean.splitlines():
            m = re.match(r"Employer Name:\s*(.+)", line.strip(), re.IGNORECASE)
            if m:
                emp_name = m.group(1).strip().lower()
                break

        # Fall back to two-line format "Employer Name:\n<value>"
        if not emp_name:
            lines = clean.splitlines()
            for i, line in enumerate(lines):
                if re.match(r"Employer Name:\s*$", line.strip(), re.IGNORECASE):
                    for j in range(i + 1, min(i + 4, len(lines))):
                        nxt = lines[j].strip()
                        if nxt:
                            emp_name = nxt.lower()
                            break
                    break

        # Look up the matching summary row by employer name
        summary_entry = emp_sum_by_name.get(emp_name)
        employment_checks.append(
            parse_employment_check(text, num, summary_entry)
        )

    # ── Parse reference detail pages ──────────────────────────────────────────
    professional_reference_checks = []
    for num, text in sorted(reference_pages.items()):
        # Map check number to summary list index (1-based → 0-based)
        summary_entry = ref_sum_list[num - 1] if num - 1 < len(ref_sum_list) else None
        professional_reference_checks.append(
            parse_reference_check(text, num, summary_entry)
        )

    # ── Parse other / database check pages ───────────────────────────────────
    other_checks = parse_other_checks(other_check_pages)

    # Cross-reference summary verdicts into the other_checks list.
    # Detail pages only describe the check source/methodology; the actual
    # verdict text (Cleared / Not Cleared) lives on the summary page.
    summary_db = {
        entry["check"].lower(): entry
        for entry in summary.get("database_check_summary", [])
    }
    for oc in other_checks:
        key = oc["check_name"].lower()
        # Try exact match first, then partial match
        match = summary_db.get(key)
        if not match:
            for k, v in summary_db.items():
                if key in k or k in key:
                    match = v
                    break
        if match:
            # Overwrite the detail-page result/status with the summary verdict
            oc["result"] = match["result"]
            oc["status"] = match["status"]

    # ── Assemble and return the final structured document ─────────────────────
    return {
        "source_file":                   file_name,
        "extracted_at":                  datetime.now().isoformat(timespec="seconds"),
        "total_pages":                   len(pages),
        "report_summary":                summary,
        "employment_checks":             employment_checks,
        "professional_reference_checks": professional_reference_checks,
        "other_checks":                  other_checks,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Excel / Word styling constants
# ─────────────────────────────────────────────────────────────────────────────
_DARK_BLUE  = "1F3864"   # used for section header row backgrounds
_MID_BLUE   = "2E75B6"   # used for column header row backgrounds
_LIGHT_BLUE = "DEEAF1"   # used for alternating data row tint
_WHITE      = "FFFFFF"
_BLACK      = "000000"

# Standard thin border applied to all table cells
_THIN        = Side(style="thin", color=_BLACK)
_BORDER_ALL  = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_BORDER_NONE = Border()   # no border (used for non-table areas)


# ─────────────────────────────────────────────────────────────────────────────
# Excel cell helpers — centralise styling so export functions stay readable
# ─────────────────────────────────────────────────────────────────────────────
def _xl_section_header(ws, row: int, title: str, span: int) -> None:
    """
    Write a dark-blue full-width section title into the first cell of *row*
    and merge across *span* columns.
    """
    cell = ws.cell(row=row, column=1, value=title)
    cell.font      = Font(bold=True, color=_WHITE, size=11)
    cell.fill      = PatternFill("solid", fgColor=_DARK_BLUE)
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=False)
    cell.border    = _BORDER_ALL
    if span > 1:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)


def _xl_col_headers(ws, row: int, headers: list[str]) -> None:
    """Write a mid-blue bold column header row starting at the given row."""
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col, value=h)
        cell.font      = Font(bold=True, color=_WHITE, size=10)
        cell.fill      = PatternFill("solid", fgColor=_MID_BLUE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border    = _BORDER_ALL


def _xl_data_row(ws, row: int, values: list, tint: bool = False) -> None:
    """
    Write a data row. Applies a light-blue tint on odd-indexed rows (tint=True)
    for alternating row colouring. Empty strings are replaced with an em-dash.
    """
    fill = (PatternFill("solid", fgColor=_LIGHT_BLUE)
            if tint else PatternFill("solid", fgColor=_WHITE))
    for col, val in enumerate(values, start=1):
        cell = ws.cell(row=row, column=col, value=val if val != "" else "\u2013")
        cell.font      = Font(color=_BLACK, size=10)
        cell.fill      = fill
        cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        cell.border    = _BORDER_ALL


def _xl_set_col_widths(ws, widths: list[int]) -> None:
    """Set column widths (in character units) for columns 1..N of *ws*."""
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


# ─────────────────────────────────────────────────────────────────────────────
# Output path resolver — handles versioning of export files
# ─────────────────────────────────────────────────────────────────────────────
def resolve_output_path(base_dir: Path, ref_number: str, ext: str) -> Path:
    """
    Build a versioned output path:
      base_dir / <ref_number> / <ref_number>[_vN]<ext>

    Rules:
      • Creates the ref_number subfolder if it doesn't exist.
      • If <ref_number><ext> does not yet exist, uses it (v1, implicit).
      • If it exists, increments: _v2, _v3, … until a free slot is found.
      • ref_number values that are blank or contain illegal filename characters
        are sanitised; blanks fall back to "UNKNOWN_REF".

    NOTE: overwrite_existing_exports=false is enforced by always calling this
    function with overwrite=False from the UI. The overwrite=True branch in
    export_to_word / export_to_csv / export_to_json bypasses this function.
    """
    # Sanitise ref_number — replace characters illegal in Windows filenames
    ref       = re.sub(r'[<>:"/\\|?*]', "_", ref_number).strip() or "UNKNOWN_REF"
    subfolder = base_dir / ref
    subfolder.mkdir(parents=True, exist_ok=True)

    # Try the base filename first (no version suffix)
    candidate = subfolder / f"{ref}{ext}"
    if not candidate.exists():
        return candidate

    # Version-bump until we find a free filename
    version = 2
    while True:
        candidate = subfolder / f"{ref}_v{version}{ext}"
        if not candidate.exists():
            return candidate
        version += 1


# ─────────────────────────────────────────────────────────────────────────────
# Word document helpers
# ─────────────────────────────────────────────────────────────────────────────
def _word_field(doc, label: str, value: str) -> None:
    """
    Add a single paragraph with bold "Label:" followed by normal " Value".
    Empty values are replaced with an em-dash placeholder.
    """
    p    = doc.add_paragraph()
    run1 = p.add_run(f"{label}:  ")
    run1.bold       = True
    run1.font.size  = Pt(10)
    run2 = p.add_run(value if value else "\u2013")  # em-dash for empty
    run2.bold       = False
    run2.font.size  = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def _word_section(doc, title: str) -> None:
    """
    Add a blue shaded section heading paragraph.
    Uses low-level OOXML to set the paragraph shading because python-docx
    does not expose paragraph background colour through its high-level API.
    """
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement

    p   = doc.add_paragraph()
    run = p.add_run(f"  {title}  ")
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)   # white text

    # Apply blue background via OOXML shading element
    shading = p._element.get_or_add_pPr()
    shd     = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "2E75B6")    # medium blue
    shading.append(shd)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)


# ─────────────────────────────────────────────────────────────────────────────
# Export: Word (.docx)
# ─────────────────────────────────────────────────────────────────────────────
def export_to_word(file_name: str, structured: dict,
                   ref_number: str, overwrite: bool) -> Path:
    """
    Save the structured report data to a formatted .docx file.

    Output location:
      WORD_OUT_DIR / <ref_number> / <ref_number>[_vN].docx

    overwrite=False (default): version-bump if the file already exists.
    overwrite=True:  always write to <ref_number>.docx, replacing it in-place.

    Returns the Path of the written file.
    """
    if overwrite:
        # Overwrite mode: always use the base filename, no version suffix
        ref       = re.sub(r'[<>:"/\\|?*]', "_", ref_number).strip() or "UNKNOWN_REF"
        subfolder = WORD_OUT_DIR / ref
        subfolder.mkdir(parents=True, exist_ok=True)
        out_path  = subfolder / f"{ref}.docx"
    else:
        # Version-bump mode: find the next available _vN filename
        out_path = resolve_output_path(WORD_OUT_DIR, ref_number, ".docx")

    doc     = Document()
    summary = structured.get("report_summary", {})

    # ── Document title block ──────────────────────────────────────────────────
    title = doc.add_heading("PDF Background Check Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"File: {file_name}    |    "
        f"Extracted: {datetime.now().strftime('%d %b %Y  %H:%M')}"
    ).paragraph_format.space_after = Pt(10)

    # ── Report summary fields ─────────────────────────────────────────────────
    _word_section(doc, "REPORT SUMMARY")
    _word_field(doc, "Subject Name",   summary.get("subject_name",  ""))
    _word_field(doc, "Overall Status", summary.get("overall_status",""))
    _word_field(doc, "Case Reference", summary.get("case_reference",""))
    _word_field(doc, "Case Received",  summary.get("case_received", ""))
    _word_field(doc, "Package",        summary.get("package",       ""))
    _word_field(doc, "Delivery Date",  summary.get("delivery_date", ""))

    # ── Employment check summary table ────────────────────────────────────────
    emp_sum = summary.get("employment_check_summary", [])
    if emp_sum:
        doc.add_paragraph()
        _word_section(doc, "EMPLOYMENT CHECK SUMMARY")
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Employer", "Result", "Status"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for e in emp_sum:
            row = tbl.add_row().cells
            row[0].text = e.get("employer", "")
            row[1].text = e.get("result",   "")
            row[2].text = e.get("status",   "")

    # ── Database check summary table ──────────────────────────────────────────
    db_sum = summary.get("database_check_summary", [])
    if db_sum:
        doc.add_paragraph()
        _word_section(doc, "DATABASE CHECK SUMMARY")
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Check", "Result", "Status"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for d in db_sum:
            row = tbl.add_row().cells
            row[0].text = d.get("check",  "")
            row[1].text = d.get("result", "")
            row[2].text = d.get("status", "")

    # ── Employment check detail sections ─────────────────────────────────────
    for ec in structured.get("employment_checks", []):
        doc.add_paragraph()
        _word_section(doc, f"EMPLOYMENT CHECK ({ec['check_number']})")
        _word_field(doc, "Employer Name",       ec.get("employer_name",        ""))
        _word_field(doc, "Company Address",      ec.get("company_address",      ""))
        _word_field(doc, "Position Title",       ec.get("position_title",       ""))
        _word_field(doc, "Dates of Employment",  ec.get("dates_of_employment",  ""))
        _word_field(doc, "Status of Employment", ec.get("status_of_employment", ""))
        _word_field(doc, "Eligible for Rehire",  ec.get("eligible_for_rehire",  ""))
        _word_field(doc, "Reason for Exit",      ec.get("reason_for_exit",      ""))
        _word_field(doc, "Respondent\u2019s Name",  ec.get("respondents_name",  ""))
        _word_field(doc, "Respondent\u2019s Title", ec.get("respondents_title", ""))
        _word_field(doc, "Contact Details",      ec.get("contact_details",      ""))
        _word_field(doc, "Verification Date",    ec.get("verification_date",    ""))
        _word_field(doc, "Notes",                ec.get("notes",                ""))
        _word_field(doc, "Verification Status",  ec.get("verification_status",  ""))

    # ── Professional reference check detail sections ───────────────────────────
    for rc in structured.get("professional_reference_checks", []):
        doc.add_paragraph()
        _word_section(doc, f"PROFESSIONAL REFERENCE CHECK ({rc['check_number']})")
        _word_field(doc, "Referee Name",                  rc.get("referee_name",        ""))
        _word_field(doc, "Result",                        rc.get("result",               ""))
        _word_field(doc, "Verification Status",           rc.get("verification_status",  ""))
        _word_field(doc, "Verifier\u2019s Name & Designation", rc.get("verifiers_name", ""))
        _word_field(doc, "Verifier\u2019s Contact",       rc.get("verifiers_contact",    ""))
        _word_field(doc, "Notes",                         rc.get("notes",                ""))
        # Write each Q&A pair
        for qa in rc.get("qa", []):
            if qa.get("question"):
                doc.add_paragraph()
                _word_field(doc, "Q", qa["question"])
                _word_field(doc, "A", qa.get("answer", ""))

    # ── Other / database check detail sections ────────────────────────────────
    if structured.get("other_checks"):
        doc.add_paragraph()
        _word_section(doc, "OTHER CHECKS")
        for oc in structured["other_checks"]:
            doc.add_paragraph()
            _word_field(doc, "Check",  oc.get("check_name", ""))
            _word_field(doc, "Result", oc.get("result",     ""))
            _word_field(doc, "Status", oc.get("status",     ""))
            _word_field(doc, "Source", oc.get("source",     ""))

    doc.save(str(out_path))
    log.info("WORD  → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Export: Excel (.xlsx)
# Function is named export_to_csv for historical reasons but produces .xlsx.
# ─────────────────────────────────────────────────────────────────────────────
def export_to_csv(file_name: str, structured: dict,
                  ref_number: str, overwrite: bool) -> Path:
    """
    Save structured data to a formatted Excel workbook (.xlsx).

    Output location:
      CSV_OUT_DIR / <ref_number> / <ref_number>[_vN].xlsx

    Sheets:
      1. Report Summary   — header fields + employment/DB summary sub-tables
      2. Employment Checks — one row per check (if present)
      3. Reference Checks  — one row per referee (if present)
      4. Other Checks      — one row per database check (if present)

    Returns the Path of the written file.
    """
    if overwrite:
        ref       = re.sub(r'[<>:"/\\|?*]', "_", ref_number).strip() or "UNKNOWN_REF"
        subfolder = CSV_OUT_DIR / ref
        subfolder.mkdir(parents=True, exist_ok=True)
        out_path  = subfolder / f"{ref}.xlsx"
    else:
        out_path = resolve_output_path(CSV_OUT_DIR, ref_number, ".xlsx")

    wb      = Workbook()
    summary = structured.get("report_summary", {})

    # ── Sheet 1: Report Summary ───────────────────────────────────────────────
    ws       = wb.active
    ws.title = "Report Summary"
    ws.freeze_panes = "A3"   # freeze header rows so they stay visible when scrolling

    _xl_section_header(ws, 1, "REPORT SUMMARY", 2)

    # Core header fields — label in col A, value in col B
    fields = [
        ("Subject Name",   summary.get("subject_name",  "")),
        ("Overall Status", summary.get("overall_status","")),
        ("Case Reference", summary.get("case_reference","")),
        ("Case Received",  summary.get("case_received", "")),
        ("Package",        summary.get("package",       "")),
        ("Delivery Date",  summary.get("delivery_date", "")),
    ]
    for r, (label, value) in enumerate(fields, start=2):
        lc = ws.cell(row=r, column=1, value=label)
        lc.font      = Font(bold=True, size=10)
        lc.fill      = PatternFill("solid", fgColor=_LIGHT_BLUE)
        lc.alignment = Alignment(horizontal="left", vertical="center")
        lc.border    = _BORDER_ALL
        vc = ws.cell(row=r, column=2, value=value if value else "\u2013")
        vc.font      = Font(size=10)
        vc.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        vc.border    = _BORDER_ALL
    _xl_set_col_widths(ws, [28, 40])

    # Employment check summary sub-table (appended below the header fields)
    emp_sum = summary.get("employment_check_summary", [])
    if emp_sum:
        start_r = 2 + len(fields) + 2   # leave one blank row as spacer
        _xl_section_header(ws, start_r, "EMPLOYMENT CHECK SUMMARY", 4)
        _xl_col_headers(ws, start_r + 1, ["Employer", "Result", "Status", "Note"])
        for i, e in enumerate(emp_sum):
            _xl_data_row(ws, start_r + 2 + i,
                         [e.get("employer",""), e.get("result",""),
                          e.get("status",""),   e.get("note","")],
                         tint=(i % 2 == 1))

    # Database check summary sub-table (appended below the employment table)
    db_sum = summary.get("database_check_summary", [])
    if db_sum:
        # Position the DB table below the employment table (or directly below fields)
        start_r2 = (start_r + 2 + len(emp_sum) + 2) if emp_sum else (2 + len(fields) + 2)
        _xl_section_header(ws, start_r2, "DATABASE CHECK SUMMARY", 3)
        _xl_col_headers(ws, start_r2 + 1, ["Check", "Result", "Status"])
        for i, d in enumerate(db_sum):
            _xl_data_row(ws, start_r2 + 2 + i,
                         [d.get("check",""), d.get("result",""), d.get("status","")],
                         tint=(i % 2 == 1))

    ws.row_dimensions[1].height = 20   # slightly taller section header row

    # ── Sheet 2: Employment Checks ────────────────────────────────────────────
    emp_checks = structured.get("employment_checks", [])
    if emp_checks:
        ws2              = wb.create_sheet("Employment Checks")
        ws2.freeze_panes = "A3"
        _xl_section_header(ws2, 1, "EMPLOYMENT CHECKS", 14)
        _xl_col_headers(ws2, 2, [
            "Check #", "Employer Name", "Company Address", "Position Title",
            "Dates of Employment", "Status of Employment", "Eligible for Rehire",
            "Reason for Exit", "Respondent\u2019s Name", "Respondent\u2019s Title",
            "Contact Details", "Verification Date", "Notes", "Verification Status",
        ])
        for i, ec in enumerate(emp_checks):
            _xl_data_row(ws2, 3 + i, [
                ec.get("check_number",        ""),
                ec.get("employer_name",        ""),
                ec.get("company_address",      ""),
                ec.get("position_title",       ""),
                ec.get("dates_of_employment",  ""),
                ec.get("status_of_employment", ""),
                ec.get("eligible_for_rehire",  ""),
                ec.get("reason_for_exit",      ""),
                ec.get("respondents_name",     ""),
                ec.get("respondents_title",    ""),
                ec.get("contact_details",      ""),
                ec.get("verification_date",    ""),
                ec.get("notes",                ""),
                ec.get("verification_status",  ""),
            ], tint=(i % 2 == 1))
        _xl_set_col_widths(ws2, [8, 22, 22, 20, 24, 20, 18, 16, 20, 20, 28, 18, 20, 18])
        ws2.row_dimensions[1].height = 20
        ws2.row_dimensions[2].height = 30   # taller column header row

    # ── Sheet 3: Professional Reference Checks ────────────────────────────────
    ref_checks = structured.get("professional_reference_checks", [])
    if ref_checks:
        ws3              = wb.create_sheet("Reference Checks")
        ws3.freeze_panes = "A3"
        _xl_section_header(ws3, 1, "PROFESSIONAL REFERENCE CHECKS", 7)
        _xl_col_headers(ws3, 2, [
            "Check #", "Referee Name", "Result", "Verification Status",
            "Verifier\u2019s Name & Designation", "Verifier\u2019s Contact", "Notes",
        ])
        for i, rc in enumerate(ref_checks):
            _xl_data_row(ws3, 3 + i, [
                rc.get("check_number",       ""),
                rc.get("referee_name",        ""),
                rc.get("result",              ""),
                rc.get("verification_status", ""),
                rc.get("verifiers_name",      ""),
                rc.get("verifiers_contact",   ""),
                rc.get("notes",               ""),
            ], tint=(i % 2 == 1))
        _xl_set_col_widths(ws3, [8, 22, 28, 18, 30, 28, 14])
        ws3.row_dimensions[1].height = 20
        ws3.row_dimensions[2].height = 30

    # ── Sheet 4: Other / Database Checks ─────────────────────────────────────
    other = structured.get("other_checks", [])
    if other:
        ws4              = wb.create_sheet("Other Checks")
        ws4.freeze_panes = "A3"
        _xl_section_header(ws4, 1, "OTHER CHECKS", 4)
        _xl_col_headers(ws4, 2, ["Check Name", "Result", "Status", "Source"])
        for i, oc in enumerate(other):
            _xl_data_row(ws4, 3 + i, [
                oc.get("check_name", ""),
                oc.get("result",     ""),
                oc.get("status",     ""),
                oc.get("source",     ""),
            ], tint=(i % 2 == 1))
        _xl_set_col_widths(ws4, [30, 38, 14, 60])
        ws4.row_dimensions[1].height = 20
        ws4.row_dimensions[2].height = 30

    wb.save(str(out_path))
    log.info("XLSX  → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Export: JSON
# ─────────────────────────────────────────────────────────────────────────────
def export_to_json(file_name: str, structured: dict,
                   ref_number: str, overwrite: bool) -> Path:
    """
    Save the structured extraction result to a pretty-printed JSON file.

    Output location:
      JSON_OUT_DIR / <ref_number> / <ref_number>[_vN].json

    Returns the Path of the written file.
    """
    if overwrite:
        ref       = re.sub(r'[<>:"/\\|?*]', "_", ref_number).strip() or "UNKNOWN_REF"
        subfolder = JSON_OUT_DIR / ref
        subfolder.mkdir(parents=True, exist_ok=True)
        out_path  = subfolder / f"{ref}.json"
    else:
        out_path = resolve_output_path(JSON_OUT_DIR, ref_number, ".json")

    with open(out_path, "w", encoding="utf-8") as fh:
        json.dump(structured, fh, indent=2, ensure_ascii=False)

    log.info("JSON  → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Full pipeline — processes a single Box PDF end-to-end
# Used when running this script directly (not via the UI).
# ─────────────────────────────────────────────────────────────────────────────
def process_pdf(client: Client, box_file: dict,
                password: str, overwrite: bool) -> None:
    """
    Full pipeline for a single Box PDF:
      download → decrypt → extract text → parse → export (Word + Excel + JSON)

    box_file — a {id, name} dict as returned by find_pdf_files_on_box().
    """
    file_id   = box_file["id"]
    file_name = box_file["name"]

    log.info("=" * 60)
    log.info("Processing: %s (Box id=%s)", file_name, file_id)

    # ── Download ──────────────────────────────────────────────────────────────
    try:
        pdf_bytes = download_pdf_bytes(client, file_id, file_name)
    except Exception as exc:
        log.error("Failed to download '%s': %s", file_name, exc)
        return

    # ── Decrypt + extract text ────────────────────────────────────────────────
    try:
        doc   = open_and_decrypt_pdf(pdf_bytes, file_name, password)
        pages = extract_text_by_page(doc)
        doc.close()
    except (ValueError, fitz.FileDataError) as exc:
        log.error("Could not read '%s': %s", file_name, exc)
        return

    log.info("Extracted %d page(s).", len(pages))

    # ── Parse into structured data ────────────────────────────────────────────
    structured = build_structured_json(file_name, pages)
    emp_count  = len(structured.get("employment_checks", []))
    ref_count  = len(structured.get("professional_reference_checks", []))
    oth_count  = len(structured.get("other_checks", []))
    log.info("Parsed: %d employment, %d reference, %d other check(s).",
             emp_count, ref_count, oth_count)

    # Use the Case Reference No from the parsed data as the output folder/filename.
    # Falls back to the PDF filename stem if the reference number is not found.
    ref_number = structured.get("report_summary", {}).get("case_reference", "").strip()
    if not ref_number:
        ref_number = Path(file_name).stem
    log.info("Reference number: %s", ref_number)

    # ── Export ────────────────────────────────────────────────────────────────
    export_to_word(file_name, structured, ref_number, overwrite)
    export_to_csv( file_name, structured, ref_number, overwrite)
    export_to_json(file_name, structured, ref_number, overwrite)


# ─────────────────────────────────────────────────────────────────────────────
# Entry point — CLI / direct execution
# ─────────────────────────────────────────────────────────────────────────────
def main() -> None:
    """
    Run the full extraction pipeline from the command line.
    Reads configuration from config.json, authenticates with Box, finds all
    PDFs in the configured folder, and processes each one.
    Not invoked when the UI (pdf_extractor_ui.py) drives the pipeline.
    """
    log.info("PDF Text Extractor — starting")

    try:
        cfg = load_config()
    except FileNotFoundError as exc:
        log.error(str(exc))
        return

    password   = cfg.get("pdf_password", "")
    box_cfg    = cfg.get("box", {})
    folder_id  = box_cfg.get("folder_id", "0")
    search_sub = cfg.get("settings", {}).get("search_subfolders", True)
    overwrite  = cfg.get("settings", {}).get("overwrite_existing_exports", False)

    # Ensure output directories exist before starting
    for out_dir in (WORD_OUT_DIR, CSV_OUT_DIR, JSON_OUT_DIR):
        out_dir.mkdir(parents=True, exist_ok=True)

    # Authenticate and find PDFs
    try:
        client = get_box_client(box_cfg)
    except ValueError as exc:
        log.error(str(exc))
        return

    pdf_files = find_pdf_files_on_box(client, folder_id, search_sub)
    if not pdf_files:
        log.info("No PDF files found in Box folder %s.", folder_id)
        return

    # Process each PDF found
    for box_file in pdf_files:
        process_pdf(client, box_file, password, overwrite)

    log.info("PDF Text Extractor — finished")


if __name__ == "__main__":
    main()
